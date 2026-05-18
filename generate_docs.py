# -*- coding: utf-8 -*-
"""生成关联规则作业的四个 Word 文档。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def _set_default_font(doc, font_name="宋体", size=12):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_heading(doc, text, level=1, size=None):
    sizes = {1: 18, 2: 15, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 13))
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "黑体")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_para(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    from docx.oxml import OxmlElement
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        from docx.oxml import OxmlElement
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "黑体")
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            from docx.oxml import OxmlElement
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")
    return table


def _add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


# =====================================================================
# 第1题
# =====================================================================
def build_q1():
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, "第一题  Apriori 算法与 FP 树构建", level=1)

    _add_heading(doc, "题目", level=2)
    _add_para(doc,
              "已知某小型购物篮数据集（共5条记录），每条记录为用户单次下单的商品集合，数据如下："
              "T1={A,B,C}；T2={A,B,D}；T3={A,C,D}；T4={B,C,D}；T5={A,B,C,D}。")
    _add_para(doc,
              "（1）设最小支持度min_support=40%（即支持度≥2/5=0.4），使用Apriori算法挖掘所有频繁项集（L1、L2、L3、L4）；")
    _add_para(doc,
              "（2）基于上述频繁项集，构建Frequent-pattern tree（FP树），标注各节点的支持度；")
    _add_para(doc,
              "（3）简要说明Apriori算法生成频繁项集的核心逻辑，以及FP树相比Apriori算法在候选集生成上的优势。")

    _add_heading(doc, "（1）Apriori 算法挖掘频繁项集", level=2)
    _add_para(doc, "最小支持度 min_support = 40%，对应最小支持计数 = 2。")

    _add_heading(doc, "步骤1：生成 L1（频繁 1-项集）", level=3)
    _add_table(doc,
               ["项集", "支持度计数", "支持度", "是否频繁"],
               [["{A}", 4, 0.8, "是"],
                ["{B}", 4, 0.8, "是"],
                ["{C}", 4, 0.8, "是"],
                ["{D}", 4, 0.8, "是"]])
    _add_para(doc, "L1 = { {A}:4, {B}:4, {C}:4, {D}:4 }", bold=True)

    _add_heading(doc, "步骤2：由 L1 自连接生成 C2，扫描得 L2", level=3)
    _add_table(doc,
               ["项集", "支持度计数", "支持度", "是否频繁"],
               [["{A,B}", 3, 0.6, "是"],
                ["{A,C}", 3, 0.6, "是"],
                ["{A,D}", 3, 0.6, "是"],
                ["{B,C}", 3, 0.6, "是"],
                ["{B,D}", 3, 0.6, "是"],
                ["{C,D}", 3, 0.6, "是"]])
    _add_para(doc, "L2 = { {A,B}:3, {A,C}:3, {A,D}:3, {B,C}:3, {B,D}:3, {C,D}:3 }", bold=True)

    _add_heading(doc, "步骤3：由 L2 自连接生成 C3，剪枝并扫描得 L3", level=3)
    _add_table(doc,
               ["项集", "支持度计数", "支持度", "是否频繁"],
               [["{A,B,C}", 2, 0.4, "是"],
                ["{A,B,D}", 2, 0.4, "是"],
                ["{A,C,D}", 2, 0.4, "是"],
                ["{B,C,D}", 2, 0.4, "是"]])
    _add_para(doc, "L3 = { {A,B,C}:2, {A,B,D}:2, {A,C,D}:2, {B,C,D}:2 }", bold=True)

    _add_heading(doc, "步骤4：由 L3 生成 C4", level=3)
    _add_table(doc,
               ["项集", "支持度计数", "支持度", "是否频繁"],
               [["{A,B,C,D}", 1, 0.2, "否"]])
    _add_para(doc, "L4 = ∅，算法终止。最终频繁项集 = L1 ∪ L2 ∪ L3。", bold=True)

    _add_heading(doc, "（2）构建 FP 树", level=2)
    _add_para(doc, "由于 A、B、C、D 支持度均为 4，按字典序排列头表顺序为：A → B → C → D。")

    _add_heading(doc, "头表（Header Table）", level=3)
    _add_table(doc,
               ["Item", "Support Count"],
               [["A", 4], ["B", 4], ["C", 4], ["D", 4]])

    _add_heading(doc, "事务按头表顺序重排", level=3)
    _add_table(doc,
               ["原事务", "重排后"],
               [["{A,B,C}", "A,B,C"],
                ["{A,B,D}", "A,B,D"],
                ["{A,C,D}", "A,C,D"],
                ["{B,C,D}", "B,C,D"],
                ["{A,B,C,D}", "A,B,C,D"]])

    _add_heading(doc, "FP 树结构（节点:支持度）", level=3)
    fp_tree = (
        "root\n"
        "├── A:4\n"
        "│   ├── B:3\n"
        "│   │   ├── C:2\n"
        "│   │   │   └── D:1\n"
        "│   │   └── D:1\n"
        "│   └── C:1\n"
        "│       └── D:1\n"
        "└── B:1\n"
        "    └── C:1\n"
        "        └── D:1\n"
    )
    _add_code(doc, fp_tree)

    _add_para(doc, "节点链表（同名节点连接，便于挖掘）：")
    _add_bullet(doc, "A 链：A:4")
    _add_bullet(doc, "B 链：A 子节点 B:3 → root 子节点 B:1")
    _add_bullet(doc, "C 链：B 下 C:2 → A 下 C:1 → 另一分支 B 下 C:1")
    _add_bullet(doc, "D 链：上述各 C 节点下的 D:1，以及 B:3 直接子节点 D:1")

    _add_heading(doc, "（3）核心逻辑与优势对比", level=2)
    _add_para(doc, "Apriori 算法的核心逻辑：基于 Apriori 性质——"
                   "频繁项集的任一非空子集必为频繁项集，反之非频繁项集的所有超集必为非频繁项集。")
    _add_bullet(doc, "由 L(k-1) 自连接生成候选 Ck；")
    _add_bullet(doc, "利用反单调性进行剪枝，剔除任一 (k-1) 子集不属于 L(k-1) 的候选；")
    _add_bullet(doc, "扫描数据库统计候选支持度，保留 ≥ min_support 的项集得到 Lk；")
    _add_bullet(doc, "迭代直至 Lk 为空。")

    _add_para(doc, "FP 树相比 Apriori 在候选集生成上的优势：")
    _add_table(doc,
               ["维度", "Apriori", "FP-growth (FP 树)"],
               [["候选集", "需显式生成大量候选项集", "无需生成候选集"],
                ["数据库扫描次数", "每层扫描一次（共 k+1 次）", "仅扫描两次"],
                ["内存与 I/O", "候选爆炸，I/O 频繁", "前缀共享压缩，内存高效"],
                ["适用规模", "适合小数据/低维", "适合大数据/高维稀疏"]])
    _add_para(doc, "核心优势可概括为：无候选生成 + 数据高度压缩 + 仅两次扫库。", bold=True)

    out_path = os.path.join(OUT_DIR, "第一题.docx")
    doc.save(out_path)
    return out_path


# =====================================================================
# 第2题
# =====================================================================
def build_q2():
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, "第二题  电商场景下 Apriori 与 FP-growth 工程对比", level=1)

    _add_heading(doc, "题目", level=2)
    _add_para(doc,
              "电商平台积累了500万条用户购物篮数据（每条数据包含用户单次下单的所有商品），商品种类达800+，"
              "核心需求是挖掘商品之间的关联规则（如\"购买牛奶→购买面包\"），用于商品组合推荐，要求挖掘效率高、"
              "规则冗余少，适配电商实时推荐的工程需求。")
    _add_para(doc, "（1）分析 Apriori 算法在该工程场景中应用的核心痛点（结合数据规模和算法原理）；")
    _add_para(doc, "（2）说明 FP-growth 算法如何解决上述痛点，结合算法原理简述其在该场景的工程实现步骤；")
    _add_para(doc, "（3）对比 Apriori 算法与 FP-growth 算法在该场景的工程适用性，说明该场景应优先选择哪种算法及原因。")

    _add_heading(doc, "（1）Apriori 算法的核心痛点", level=2)
    _add_bullet(doc, "候选项集爆炸：800 种商品在 k=2 时候选已达 C(800,2)≈32 万，k 增大后呈指数级膨胀，内存压力极大。")
    _add_bullet(doc, "多次扫描磁盘 I/O 代价高：挖到 k 层频繁项集需扫描数据库 k 次，500 万条记录配合多轮 I/O，时延为分钟到小时级，无法满足实时推荐。")
    _add_bullet(doc, "支持度阈值难权衡：阈值高则规则太少漏掉长尾组合，阈值低则候选爆炸；电商商品长尾分布显著，矛盾尤其突出。")
    _add_bullet(doc, "冗余规则多：Apriori 不区分闭频繁项集与最大频繁项集，输出大量冗余规则，需额外后处理。")
    _add_bullet(doc, "增量与并行差：电商订单流式新增，Apriori 每次需重跑全量；候选连接逻辑也不易自然分布式拆分。")

    _add_heading(doc, "（2）FP-growth 解决思路与工程实现步骤", level=2)
    _add_para(doc, "原理层解决方案：")
    _add_bullet(doc, "用 FP-tree 共享前缀压缩 500 万事务，相同热门商品组合（如\"牛奶+面包\"）只需一条路径计数，存储与扫描开销骤降；")
    _add_bullet(doc, "仅需两次扫库（统计单项频度 + 建树），消除多轮 I/O；")
    _add_bullet(doc, "通过条件模式基（Conditional Pattern Base）→ 条件 FP 树递归挖掘，完全省掉候选生成步骤。")

    _add_para(doc, "电商场景工程实现步骤：")
    _add_bullet(doc, "数据预处理：从订单库抽取 user_id × order_id × item_list，过滤无效订单，按 order_id 聚合成事务集，落盘至 HDFS/Parquet。")
    _add_bullet(doc, "第一次扫描：用 MapReduce/Spark 全表统计每件商品支持计数，过滤低于 min_support 的商品，得到频繁 1-项集 F1，按降序构建头表。")
    _add_bullet(doc, "第二次扫描：将每条事务按 F1 顺序重排，并发插入 FP-tree；分布式可采用 PFP-Growth：按头表分组将事务路由到不同分区，各自构建子 FP-tree。")
    _add_bullet(doc, "挖掘阶段：对头表自底向上，针对每个频繁项 i 抽取条件模式基 → 构建条件 FP-tree → 递归挖掘，生成完整频繁项集。")
    _add_bullet(doc, "规则生成与剪枝：基于频繁项集生成规则，结合 confidence、lift、leverage 多指标过滤，保留 lift>1 且置信度达阈值的强关联规则。")
    _add_bullet(doc, "上线服务：将规则写入 Redis/HBase，键为前件商品，值为后件 Top-N 推荐列表，供推荐引擎毫秒级查询；定时（小时/天级）增量重训以跟进趋势。")

    _add_heading(doc, "（3）适用性对比与选型结论", level=2)
    _add_table(doc,
               ["对比维度", "Apriori", "FP-growth"],
               [["数据库扫描次数", "k 次", "2 次"],
                ["候选集生成", "大量显式生成", "无"],
                ["内存占用（800+商品×500万事务）", "候选集爆炸，常 OOM", "树结构高度压缩"],
                ["时间复杂度", "高（指数级）", "显著更低"],
                ["分布式扩展性", "较弱", "有成熟 PFP-Growth"],
                ["实时/准实时推荐适配", "难", "适配"],
                ["规则冗余", "多", "较少（结合闭模式更优）"]])
    _add_para(doc, "结论：该场景应优先选择 FP-growth。", bold=True)
    _add_para(doc,
              "原因：电商 500 万条订单、800+ 商品规模下，Apriori 候选爆炸与多次扫库会导致挖掘耗时不可接受；"
              "FP-growth 通过 FP-tree 压缩、仅两次扫库、无候选生成三大优势，可在分布式集群上将整轮挖掘压缩到分钟级，"
              "输出的关联规则可直接落地实时推荐引擎，满足\"高效率、低冗余、可上线\"的工程需求。")

    out_path = os.path.join(OUT_DIR, "第二题.docx")
    doc.save(out_path)
    return out_path


# =====================================================================
# 第3题
# =====================================================================
def build_q3():
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, "第三题  电商用户行为分析中关联规则挖掘的核心价值", level=1)

    _add_heading(doc, "题目", level=2)
    _add_para(doc,
              "在电商平台用户行为分析工程场景中（核心需求：挖掘用户浏览、加购、下单、支付等行为的关联关系，"
              "辅助平台制定个性化推荐策略），请分析关联规则挖掘在该场景中的核心价值是什么？"
              "并说明该场景下，关联规则挖掘与传统用户画像分析的核心区别，贴合工程实践需求。")

    _add_heading(doc, "一、关联规则挖掘的核心价值", level=2)
    _add_para(doc, "在\"浏览→加购→下单→支付\"行为链路中，关联规则挖掘的核心价值体现在四个方面：")
    _add_bullet(doc, "揭示行为链路的转化关联：挖掘出\"浏览 A 品类 → 加购 B 商品（支持度 2%、置信度 45%、lift=3.2）\"等规则，量化前置行为对后置行为的驱动强度，定位关键转化节点。")
    _add_bullet(doc, "驱动个性化推荐：基于\"加购商品 X → 下单商品 Y\"的强规则，构建\"凑单推荐\"\"猜你喜欢\"等场景，提升 GMV 与客单价。")
    _add_bullet(doc, "优化运营触达策略：识别\"浏览未下单\"高频路径，针对性推送优惠券/提醒以降低流失；识别\"支付即复购\"组合，推动会员复购体系。")
    _add_bullet(doc, "辅助商品与活动设计：发现高关联商品组合，指导套餐打包、连带陈列、促销活动设计（如满减门槛设定）；规则的时序版本（序列模式）还可揭示行为趋势。")

    _add_heading(doc, "二、与传统用户画像分析的核心区别", level=2)
    _add_table(doc,
               ["对比维度", "传统用户画像", "关联规则挖掘"],
               [["分析对象", "单个用户的静态属性（性别、年龄、消费等级、偏好标签）",
                 "行为之间的关系（行为A → 行为B 的关联）"],
                ["输出形态", "用户标签集合，如\"女、25 岁、母婴偏好、高消费\"",
                 "规则集合，如\"浏览奶粉 → 加购纸尿裤，置信度 58%\""],
                ["数据视角", "以\"人\"为中心的纵向汇总", "以\"事务/会话\"为中心的横向共现"],
                ["关注点", "用户是谁（Who）", "行为怎么发生（How / Why together）"],
                ["时效性", "偏长期、缓变", "偏中短期、可快速迭代"],
                ["推荐机制", "基于相似用户/标签匹配", "基于规则前后件直接推断推荐内容"],
                ["工程产物", "标签库、用户分群", "规则库、推荐召回策略"],
                ["解释性", "描述用户特征", "可解释的因果/共现关系，业务方易理解"]])

    _add_para(doc,
              "工程实践中两者互补：用户画像负责\"圈定推荐对象\"，关联规则负责\"决定推荐内容\"。"
              "先用画像分群（如母婴高活用户），再在群内应用关联规则（\"加购奶粉→推荐纸尿裤\"），"
              "两者结合可显著提升推荐精度与转化率。")

    out_path = os.path.join(OUT_DIR, "第三题.docx")
    doc.save(out_path)
    return out_path


# =====================================================================
# 第4题
# =====================================================================
def build_q4():
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, "第四题  金融风控中的多维关联规则挖掘", level=1)

    _add_heading(doc, "题目", level=2)
    _add_para(doc,
              "在金融风控工程场景中（核心需求：挖掘用户信贷数据、消费数据、征信数据的多维关联关系，"
              "识别信贷风险隐患），需进行多维关联规则挖掘。请分析多维关联规则挖掘与传统单维关联规则挖掘"
              "（如仅挖掘商品关联）的核心区别；结合该场景，列举3个核心维度，并说明如何通过多维关联规则挖掘"
              "识别信贷风险，给出具体的工程实现思路。")

    _add_heading(doc, "一、多维与单维关联规则的核心区别", level=2)
    _add_table(doc,
               ["维度", "单维关联规则", "多维关联规则"],
               [["涉及谓词", "同一谓词内的项关联（如 buys(X,牛奶)→buys(X,面包)）",
                 "多个不同谓词/属性间的关联（如 age, income, credit_history）"],
                ["数据形式", "单一事务表（项的集合）", "多张表/多属性的关系数据（信贷+消费+征信）"],
                ["项的语义", "项无明确属性归属", "每个项需带\"维度=值\"标签，如 income=低"],
                ["处理难点", "频繁项集爆炸", "数值属性离散化、维度组合爆炸、跨源数据融合"],
                ["算法基础", "Apriori / FP-growth 直接适用",
                 "需先离散化/概念分层，再扩展到属性-值对项空间"],
                ["典型应用", "购物篮分析", "风控、医疗诊断、用户精细画像"],
                ["输出可解释性", "\"买A的也买B\"",
                 "\"低收入∧多头借贷∧征信查询频繁→高违约风险\"，业务含义更强"]])
    _add_para(doc,
              "关键扩展：多维关联将\"项\"从\"商品\"扩展到\"<属性,值>\"二元组，因此需引入"
              "离散化（等宽/等频/聚类分箱）与概念层次（如 income：低/中/高），并区分维内、维间、混合维关联规则。")

    _add_heading(doc, "二、金融风控场景的 3 个核心维度", level=2)
    _add_table(doc,
               ["维度", "典型属性（已离散化）"],
               [["D1 信贷维度",
                 "信贷申请次数（高/中/低）、未结清贷款笔数、当前逾期天数、负债收入比 DTI（高/低）"],
                ["D2 消费维度",
                 "月均消费额（高/中/低）、夜间消费占比、博彩/高风险类商户占比、消费波动率"],
                ["D3 征信/行为维度",
                 "近 3 月征信查询次数（高/低）、多头借贷数、地址/手机号变更频率、社交关联中黑名单数"]])

    _add_heading(doc, "三、识别信贷风险的工程实现思路", level=2)
    _add_para(doc, "目标：挖掘形如下述高质量风险规则，用于评分卡补充与拒绝策略：", bold=True)
    _add_code(doc,
              "DTI=高 ∧ 近3月征信查询=高 ∧ 夜间消费占比=高 → 违约=是\n"
              "(support=3%, confidence=72%, lift=4.5)")

    _add_para(doc, "实施步骤：", bold=True)
    _add_bullet(doc, "多源数据融合：从信贷系统、交易流水、人行/三方征信抽取宽表，以 user_id 为主键 join；处理缺失/异常；统一观察期（6 个月）与表现期（3 个月）；打上目标标签 default ∈ {是, 否}（如 M3+ 逾期）。")
    _add_bullet(doc, "属性离散化与概念分层：数值属性（DTI、消费额等）使用等频/卡方/决策树分箱，确保各箱样本量均衡；类别属性建立层次（行业→大类→小类）；转换为<维度=值>二元组事务，例如{DTI=高, 查询次数=高, 夜间消费=高, 多头借贷=有, default=是}。")
    _add_bullet(doc, "频繁项集挖掘：使用 FP-growth（Spark MLlib FPGrowth）在分布式环境挖掘，min_support 视样本不平衡情况设定（如 0.5%–2%）；负样本稀少时可对其上采样或先降低阈值再过滤。")
    _add_bullet(doc, "关联规则生成与风险规则筛选：仅保留后件包含 default=是 的规则；多指标过滤 confidence≥0.6、lift≥2，并辅以 leverage/conviction；通过极大或闭频繁项集去冗余。")
    _add_bullet(doc, "业务校验与上线：风控专家审核规则的业务可解释性；在历史样本上回测 KS、坏账捕获率与覆盖度；上线为规则引擎（Drools 或自研），与评分卡、机器学习模型并联——命中高 lift 强规则直接拒绝/转人工，命中中风险规则降额提价，未命中走评分卡常规流程。")
    _add_bullet(doc, "持续迭代：监控规则在新样本上的 PSI 与 lift 衰减，按月级重训；结合在线特征仓库实现近实时更新；与图关联（团伙欺诈）、序列模式（行为时序）结合，形成多维关联+图+序列的立体风控体系。")

    _add_para(doc, "示例可挖掘规则（示意）：", bold=True)
    _add_bullet(doc, "多头借贷=有 ∧ 征信查询=高 → default=是 (conf=68%, lift=4.1)：多头风险")
    _add_bullet(doc, "DTI=高 ∧ 夜间消费占比=高 ∧ age=青年 → default=是 (conf=61%, lift=3.5)：过度消费风险")
    _add_bullet(doc, "地址变更=频繁 ∧ 手机号变更=频繁 ∧ 设备多账号=是 → default=是 (conf=75%, lift=5.8)：疑似欺诈")

    _add_para(doc,
              "通过上述多维关联规则挖掘，可在传统评分卡之外补充强可解释、易落规则引擎的风险信号，"
              "提升信贷审批的精准度与抗欺诈能力。")

    out_path = os.path.join(OUT_DIR, "第四题.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    paths = [build_q1(), build_q2(), build_q3(), build_q4()]
    for p in paths:
        print("Saved:", p)
