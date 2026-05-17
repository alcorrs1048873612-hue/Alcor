"""按"挑战杯"吉林省创业计划竞赛模板要求，将 bp.md 转换为正式 docx。"""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

SRC = 'bp.md'
DST = '吉科核源_通用小型静电高压发生器及控制系统_挑战杯创业计划书.docx'


def set_cn_font(run, cn_name='宋体', en_name='Times New Roman'):
    run.font.name = en_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), cn_name)
    rfonts.set(qn('w:ascii'), en_name)
    rfonts.set(qn('w:hAnsi'), en_name)


def add_paragraph(doc, text, style='body'):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if style == 'h1':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(16)  # 三号
        run.bold = False
        set_cn_font(run, '黑体', 'Times New Roman')
    elif style == 'h2':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(16)  # 三号
        set_cn_font(run, '楷体', 'Times New Roman')
    elif style == 'h3':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(14)  # 四号
        set_cn_font(run, '宋体', 'Times New Roman')
    elif style == 'tablename':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(12)  # 小四
        set_cn_font(run, '黑体', 'Times New Roman')
    else:  # body
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(14)  # 四号
        set_cn_font(run, '宋体', 'Times New Roman')
    return p


def add_md_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text.strip())
        run.font.size = Pt(14)
        run.bold = True
        set_cn_font(run, '宋体', 'Times New Roman')
    # 数据
    for r, row in enumerate(rows, start=1):
        for c, cell_text in enumerate(row):
            if c >= len(header):
                continue
            cell = table.rows[r].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(14)
            set_cn_font(run, '宋体', 'Times New Roman')


def setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.0)
    # 页码：在页脚居中
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(12)
    set_cn_font(run, 'Times New Roman', 'Times New Roman')
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


def parse_table_block(lines, idx):
    """解析 Markdown 表格，返回 (header, rows, next_idx)。"""
    header = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
    # 跳过分隔行
    next_idx = idx + 2
    rows = []
    while next_idx < len(lines) and lines[next_idx].lstrip().startswith('|'):
        cells = [c.strip() for c in lines[next_idx].strip().strip('|').split('|')]
        rows.append(cells)
        next_idx += 1
    return header, rows, next_idx


def main():
    doc = Document()
    setup_page(doc)

    # 标题与署名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026 年"挑战杯"吉林省大学生创业计划竞赛')
    run.font.size = Pt(18)
    set_cn_font(run, '黑体', 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('通用小型静电高压发生器及控制系统')
    run.font.size = Pt(22)
    set_cn_font(run, '黑体', 'Times New Roman')

    for line in [
        '公司名称：吉科核源科技有限公司',
        '项目组别 / 赛道：先进制造（智能制造）',
        '项目负责人：刘海远  电话：18052297885',
        '指导教师：（正高级实验师）',
        '提交日期：2026 年',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        set_cn_font(run, '宋体', 'Times New Roman')

    doc.add_page_break()

    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 跳过元信息和分隔线
        if stripped == '---':
            i += 1
            continue

        # 文档大标题（只有一行 #），跳过
        if stripped.startswith('# ') and i < 10:
            i += 1
            continue
        # 项目元信息（**项目...** 等）
        if stripped.startswith('**'):
            i += 1
            continue

        # 一级标题：# 一、xxx
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            add_paragraph(doc, text, 'h1')
            i += 1
            continue

        # 二级标题：## 1. xxx
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            add_paragraph(doc, text, 'h2')
            i += 1
            continue

        # 三级标题
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            add_paragraph(doc, text, 'h3')
            i += 1
            continue

        # 表格名（"表X.Y XXX" 单独一行）
        if re.match(r'^表\d+\.\d+', stripped):
            add_paragraph(doc, stripped, 'tablename')
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            header, rows, next_i = parse_table_block(lines, i)
            add_md_table(doc, header, rows)
            i = next_i
            continue

        # 普通段落
        add_paragraph(doc, stripped, 'body')
        i += 1

    doc.save(DST)
    print('Generated:', DST)


if __name__ == '__main__':
    main()
